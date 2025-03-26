# ---------------------------
# Standard Library Imports
# ---------------------------
import os
import re
import json
import shutil
import pickle
import logging
import asyncio
import tempfile
import warnings
import traceback
import base64
import urllib.parse
from urllib.parse import quote
import urllib3
import uuid
import io
from io import BytesIO
from datetime import datetime
from functools import partial, lru_cache
from concurrent.futures import ThreadPoolExecutor
import concurrent.futures
import win32com.client  # For handling .doc files
import pythoncom  # For COM initialization
from fake_useragent import UserAgent

# ---------------------------
# Third-Party Libraries
# ---------------------------
from dateutil.relativedelta import relativedelta
from typing import Any, Dict, List, Optional, Tuple, AsyncGenerator

# Requests & HTTP
import requests
from requests.auth import HTTPBasicAuth
from requests.packages.urllib3.exceptions import InsecureRequestWarning

# ---------------------------
# FastAPI Imports
# ---------------------------
from fastapi import (
    FastAPI,
    Query,
    HTTPException,
    File,
    UploadFile,
    Form,
    BackgroundTasks,
    WebSocket,
    WebSocketDisconnect
)
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from sse_starlette.sse import EventSourceResponse  # Added import for SSE

# ---------------------------
# Data Processing
# ---------------------------
import pandas as pd
import numpy as np
from bs4 import BeautifulSoup

# ---------------------------
# NLP & Machine Learning
# ---------------------------
import dspy
from dspy import InputField, OutputField, Signature, Predict
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
import faiss

# ---------------------------
# PDF & Document Processing
# ---------------------------
import PyPDF2
from PyPDF2 import PdfReader  # For PDF text extraction
from docx2pdf import convert
from docx import Document 
from fpdf import FPDF
from reportlab.lib.pagesizes import letter
from reportlab.platypus import (
    SimpleDocTemplate,
    BaseDocTemplate,
    Paragraph,
    Spacer,
    PageBreak,
    Table,
    TableStyle,
    Frame,
    PageTemplate
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
import reportlab.pdfgen.canvas as reportlab_canvas

# ---------------------------
# Cloud & Storage
# ---------------------------
from azure.storage.blob import BlobServiceClient

# ---------------------------
# Async Networking
# ---------------------------
import aiohttp
from aiohttp import BasicAuth, ClientTimeout

# ---------------------------
# Environment & Utility Imports
# ---------------------------
from dotenv import load_dotenv
from werkzeug.utils import secure_filename

# ---------------------------
# Pydantic Import (Added)
# ---------------------------
from pydantic import BaseModel

# Configure logging and load environment variables
logging.basicConfig(level=logging.INFO)
load_dotenv()

app = FastAPI()

# Enable CORS for local development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# UseCase 1

###########################################
# Global in-memory stores for progress messages and generated PDFs.
###########################################
progress_store: dict[str, list[str]] = {}
pdf_store: dict[str, BytesIO] = {}

class SpecInput(BaseModel):
    product_name: str
    specs: dict

# Path to the CSV file
file_path = r".\dataset.csv"

#############################################
# DSPy Setup for Extended Maintenance Insight Module
#############################################
 
class MaintenanceInsightSignature(dspy.Signature):
    """
    DSPy signature for obtaining a maintenance insight.
    The LM will receive:
    - aggregated_data: the aggregated JSON data for the motor (grouped by date)
    - question: a prompt asking for a maintenance insight, which should instruct the LM to return a JSON object
      with the following keys:
         - risk: one of "low", "moderate", or "high"
         - description: a textual description of the risk or trend (formatted as bullet points)
         - dates: a list of dates (as strings) where anomalies or risk factors were detected
         - ai_suggestions: a bullet-point formatted list of AI suggestions (including any date information as needed)
    """
    aggregated_data: str = dspy.InputField(desc="Aggregated JSON data for the motor")
    question: str = dspy.InputField(desc="Prompt for maintenance insight")
    risk: str = dspy.OutputField(desc="Risk rating: low, moderate, or high")
    description: str = dspy.OutputField(desc="Description of the risk or trend, formatted as bullet points")
    dates: str = dspy.OutputField(desc="List of relevant dates (as a JSON-formatted string)")
    ai_suggestions: str = dspy.OutputField(desc="AI-based suggestions formatted as bullet points")
 
class MaintenanceInsightModule(dspy.Module):
    def __init__(self):
        # Initialize a chain-of-thought module with the extended signature
        self.get_insight = dspy.ChainOfThought(MaintenanceInsightSignature)
   
    def forward(self, aggregated_data: str, question: str):
        result = self.get_insight(aggregated_data=aggregated_data, question=question)
        return {
            "risk": result.risk,
            "description": result.description,
            "dates": result.dates,
            "ai_suggestions": result.ai_suggestions
        }
 
# Initialize the LM for DSPy
dspy_lm = dspy.LM(
    model='azure/gpt-4o',
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_base=os.getenv("AZURE_OPENAI_ENDPOINT"),
    temperature=0.2,
    max_tokens=4096,
)
dspy.configure(lm=dspy_lm)
 
# Initialize the extended maintenance insight module
maintenance_insight_module = MaintenanceInsightModule()

#############################################
# DSPy Setup for Motor Analytics Module
#############################################

class MotorAnalyticsSignature(dspy.Signature):
    """
    DSPy signature for obtaining motor analytics insights.
    The LM will receive:
    - analytics_data: a JSON string containing analytics metrics
    - prompt: a prompt asking for AI-based observations
    It returns:
    - ai_observations: the AI-generated observations formatted accordingly.
    """
    analytics_data: str = dspy.InputField(desc="JSON string of analytics data")
    prompt: str = dspy.InputField(desc="Prompt for motor analytics insights")
    ai_observations: str = dspy.OutputField(desc="AI-based observations in specified format")

class MotorAnalyticsModule(dspy.Module):
    def __init__(self):
        self.get_analytics = dspy.ChainOfThought(MotorAnalyticsSignature)
   
    def forward(self, analytics_data: str, prompt: str):
        result = self.get_analytics(analytics_data=analytics_data, prompt=prompt)
        return result.ai_observations
 
# Initialize the motor analytics module
motor_analytics_module = MotorAnalyticsModule()
 
#############################
# Existing Endpoints
#############################
 
@app.get("/api/motor/motors")
def extract_motor_ids():
    try:
        df = pd.read_csv(file_path)
        if "Motor_ID" not in df.columns:
            raise HTTPException(status_code=400, detail="CSV file is missing the 'Motor_ID' column.")
        motor_ids = df["Motor_ID"].unique().tolist()
        return {"motors": motor_ids}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
 
@app.get("/api/motor/motor-data")
def get_motor_data(
    motor_id: str = Query(..., description="Motor ID to filter data"),
    months: int = Query(..., description="Number of months for filtering (1, 6, 12, 24, 36, 48, 60)")
):
    allowed_months = [1, 6, 12, 24, 36, 48, 60]
    if months not in allowed_months:
        raise HTTPException(status_code=400, detail=f"Invalid months parameter. Allowed values: {allowed_months}")
    try:
        df = pd.read_csv(file_path, parse_dates=["Timestamp"])
        today = datetime.today()
        start_date = today - relativedelta(months=months)
        filtered_df = df[(df["Timestamp"] >= start_date) & (df["Timestamp"] <= today) & (df["Motor_ID"] == motor_id)]
        if filtered_df.empty:
            return {"message": f"No data available for Motor ID {motor_id} in the past {months} months."}
        return {"motor_data": filtered_df.to_dict(orient="records")}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
 
#############################################
# Extended Motor Maintenance Insight Endpoint
#############################################
 
@app.get("/api/motor/motor-maintenance")
def get_motor_maintenance_insight(
    motor_id: str = Query(..., description="Motor ID for maintenance insight"),
    months: int = Query(..., description="Number of months for filtering aggregated data (1, 6, 12, 24, 36, 48, 60)")
):
    allowed_months = [1, 6, 12, 24, 36, 48, 60]
    if months not in allowed_months:
        raise HTTPException(status_code=400, detail=f"Invalid months parameter. Allowed values: {allowed_months}")
   
    try:
        df = pd.read_csv(file_path, parse_dates=["Timestamp"])
        today = datetime.today()
        start_date = today - relativedelta(months=months)
        motor_df = df[(df["Timestamp"] >= start_date) & (df["Timestamp"] <= today) & (df["Motor_ID"] == motor_id)]
        if motor_df.empty:
            return {"message": f"No data available for Motor ID {motor_id} in the past {months} months."}
       
        motor_df["Date"] = motor_df["Timestamp"].dt.date
        aggregated_df = motor_df.groupby("Date").size().reset_index(name="record_count")
        aggregated_data_json = aggregated_df.to_json(orient="records", date_format="iso", indent=2)
        logging.info(f"Aggregated data for motor {motor_id}:\n{aggregated_data_json}")
       
        anomaly_prompt = (
            f"Analyze the following aggregated maintenance data for Motor ID {motor_id} "
            f"to detect anomalies in power consumption. Identify unusual spikes or drops, list the specific dates when these anomalies occurred, "
            f"and provide AI-based suggestions for further investigation. Return your result in JSON format with keys: 'risk', 'description', 'dates', and 'ai_suggestions'. "
            f"Format both the 'description' and 'ai_suggestions' as bullet points. Data:\n{aggregated_data_json}"
        )
        anomaly_insight = maintenance_insight_module.forward(
            aggregated_data=aggregated_data_json, question=anomaly_prompt
        )
       
        risk_prompt = (
            f"Analyze the following aggregated maintenance data for Motor ID {motor_id} to determine if the motor is at high risk of failure. "
            f"Identify the dates where risk factors are evident and provide AI-based suggestions for maintenance actions. Return your result in JSON format with keys: "
            f"'risk', 'description', 'dates', and 'ai_suggestions'. Format the 'description' and 'ai_suggestions' fields as bullet points. Data:\n{aggregated_data_json}"
        )
        risk_insight = maintenance_insight_module.forward(
            aggregated_data=aggregated_data_json, question=risk_prompt
        )
       
        predictive_prompt = (
            f"Using the following aggregated maintenance data for Motor ID {motor_id}, perform AI-based predictive failure modeling. "
            f"Forecast potential future failures by identifying trends and listing the relevant dates or date ranges. Additionally, provide AI-based suggestions on how to rectify any defects observed. "
            f"Return your result in JSON format with keys: 'risk', 'description', 'dates', and 'ai_suggestions', ensuring that both 'description' and 'ai_suggestions' are formatted as bullet points. Data:\n{aggregated_data_json}"
        )
        predictive_insight = maintenance_insight_module.forward(
            aggregated_data=aggregated_data_json, question=predictive_prompt
        )
       
        if "Status" in motor_df.columns:
            status_counts = motor_df["Status"].value_counts()
        else:
            status_counts = pd.Series({"Operational": 70, "Maintenance": 20, "Failure": 10})
       
        pie_chart_data = {
            "labels": status_counts.index.tolist(),
            "datasets": [
                {
                    "label": "Motor Status",
                    "data": status_counts.tolist(),
                    "backgroundColor": ["#36A2EB", "#FFCE56", "#FF6384"],
                    "hoverBackgroundColor": ["#36A2EB", "#FFCE56", "#FF6384"]
                }
            ]
        }
       
        return {
            "motor_id": motor_id,
            "maintenance_insights": {
                "anomaly_detection": anomaly_insight,
                "high_failure_risk": risk_insight,
                "predictive_failure_modeling": predictive_insight,
            },
            "motor_status_pie_data": pie_chart_data
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
 
#############################################
# New Analytics Endpoint Using DSPy
#############################################
 
@app.get("/api/motor/analytics")
def get_motor_analytics(
    motor_id: str = Query(..., description="Motor ID for analytics"),
    months: int = Query(..., description="Number of months for filtering data (allowed: 1, 6, 12, 24, 36, 48, 60)")
):
    """
    Returns analytics data for a specific motor_id filtered by the number of months,
    and generates AI-based observations using DSPy.
    """
    allowed_months = [1, 6, 12, 24, 36, 48, 60]
    if months not in allowed_months:
        raise HTTPException(status_code=400, detail=f"Invalid months parameter. Allowed values: {allowed_months}")
 
    try:
        df = pd.read_csv(file_path)
        df.columns = df.columns.str.strip()
        df = df.rename(columns={
            "Timestamp": "timestamp",
            "Motor_ID": "motor_id",
            "Voltage (V)": "voltage",
            "Current (A)": "current",
            "Power (kW)": "power",
            "Frequency (Hz)": "frequency",
            "Power Factor": "power_factor",
            "Torque (Nm)": "torque",
            "RPM": "rpm",
            "Load (%)": "load",  
            "Temperature (Â°C)": "temperature",
            "Humidity (%)": "humidity",
            "Vibration (mm/s)": "vibration",
            "Status": "status"
        })
        df["timestamp"] = pd.to_datetime(df["timestamp"])
        df["rpm"] = pd.to_numeric(df["rpm"], errors="coerce")
        df["load"] = pd.to_numeric(df["load"], errors="coerce")
        df["vibration"] = pd.to_numeric(df["vibration"], errors="coerce")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dataset not loaded: {e}")
 
    df.columns = df.columns.str.lower()
 
    if "timestamp" not in df.columns:
        raise HTTPException(status_code=500, detail="Timestamp column not found in dataset")
    df["timestamp"] = pd.to_datetime(df["timestamp"])
    if "motor_id" not in df.columns:
        raise HTTPException(status_code=500, detail="motor_id column not found in dataset")
 
    filtered_df = df[df["motor_id"] == motor_id]
    if filtered_df.empty:
        raise HTTPException(status_code=404, detail=f"No data found for motor_id: {motor_id}")
 
    today = datetime.today()
    start_date = today - relativedelta(months=months)
    filtered_df = filtered_df[(filtered_df["timestamp"] >= start_date) & (filtered_df["timestamp"] <= today)]
    if filtered_df.empty:
        raise HTTPException(status_code=404, detail=f"No data found for motor_id: {motor_id} in the past {months} months")
 
    filtered_data = filtered_df.to_dict(orient="records")
 
    def calculate_averages(data_list, field):
        values = [float(record[field]) for record in data_list if record.get(field)]
        return sum(values) / len(values) if values else 0
     
    try:
        analytics = {
            "voltage": {
                "avg": round(calculate_averages(filtered_data, "voltage"), 2),
                "max": round(max(float(r["voltage"]) for r in filtered_data), 2),
                "min": round(min(float(r["voltage"]) for r in filtered_data), 2)
            },
            "current": {
                "avg": round(calculate_averages(filtered_data, "current"), 2),
                "max": round(max(float(r["current"]) for r in filtered_data), 2),
                "min": round(min(float(r["current"]) for r in filtered_data), 2)
            },
            "power": {
                "avg": round(calculate_averages(filtered_data, "power"), 2),
                "max": round(max(float(r["power"]) for r in filtered_data), 2),
                "min": round(min(float(r["power"]) for r in filtered_data), 2)
            },
            "load": {
                "avg": round(calculate_averages(filtered_data, "load"), 2),
                "max": round(max(float(r["load"]) for r in filtered_data), 2),
                "min": round(min(float(r["load"]) for r in filtered_data), 2)
            },
            "vibration": {
                "avg": round(calculate_averages(filtered_data, "vibration"), 2),
                "max": round(max(float(r["vibration"]) for r in filtered_data), 2),
                "min": round(min(float(r["vibration"]) for r in filtered_data), 2)
            },
            "failures": len([r for r in filtered_data if r["status"] in ["Motor Failure", "Shutdown"]]),
            "period": {
                "start": min(pd.to_datetime(r["timestamp"]) for r in filtered_data).strftime("%Y-%m-%d"),
                "end": max(pd.to_datetime(r["timestamp"]) for r in filtered_data).strftime("%Y-%m-%d")
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error computing analytics: {e}")
 
    prompt = f"""
    You are an AI trained to analyze motor performance data and provide insights. Based on the following data, generate three AI-based observations with recommendations:
    - Voltage: Avg {analytics['voltage']['avg']}V, Max {analytics['voltage']['max']}V, Min {analytics['voltage']['min']}V
    - Current: Avg {analytics['current']['avg']}A, Max {analytics['current']['max']}A, Min {analytics['current']['min']}A
    - Power: Avg {analytics['power']['avg']}W, Max {analytics['power']['max']}W, Min {analytics['power']['min']}W
    - Load: Avg {analytics['load']['avg']}%, Max {analytics['load']['max']}%, Min {analytics['load']['min']}%
    - Vibration: Avg {analytics['vibration']['avg']}, Max {analytics['vibration']['max']}, Min {analytics['vibration']['min']}
    - Failures: {analytics['failures']} incidents
    - Period: {analytics['period']['start']} to {analytics['period']['end']}
    Provide three concise insights in this format:
        Observation: [AI analysis]  
        Insight: [What's happening]  
        Recommendation: [What action to take]
    """
 
    # Use DSPy for AI observations
    analytics_data_json = json.dumps(analytics)
    ai_observations = motor_analytics_module.forward(analytics_data=analytics_data_json, prompt=prompt)
 
    return {
        "analytics": analytics,
        "ai_observations": ai_observations
    }
   
#############################################
# Trend Analysis Endpoints
#############################################
class TrendAnalysisSignature(dspy.Signature):
    """
    DSPy signature for obtaining trend analysis insights.
    The LM will receive:
    - analytics_data: a JSON string containing aggregated or trend data.
    - prompt: a prompt asking for trend analysis, instructing the LM to return a JSON object
      with keys such as 'ai_trend' that include observations and recommendations.
    """
    analytics_data: str = dspy.InputField(desc="JSON string of trend data")
    prompt: str = dspy.InputField(desc="Prompt for trend analysis insights")
    ai_trend: str = dspy.OutputField(desc="AI-generated trend insights with recommendations")
 
class TrendAnalysisModule(dspy.Module):
    def __init__(self):
        self.get_trend = dspy.ChainOfThought(TrendAnalysisSignature)
   
    def forward(self, analytics_data: str, prompt: str):
        result = self.get_trend(analytics_data=analytics_data, prompt=prompt)
        return result.ai_trend
   
trend_analysis_module = TrendAnalysisModule()
 
@app.get("/api/motor/motor-ids")
def get_motor_ids():
    """Returns list of unique motor IDs from the dataset."""
    try:
        df = pd.read_csv(file_path)
        unique_ids = df["Motor_ID"].unique().tolist()
        return {"data": unique_ids}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/motor/failure-trends")
def get_failure_trends(
    motor_id: str = Query(..., description="Motor ID for filtering data"),
    months: int = Query(..., description="Number of months for filtering (1, 6, 12, 24, 36, 48, 60)")
):
    """Returns failure trends for a specific motor over time with AI-generated trend insights."""
    allowed_months = [1, 6, 12, 24, 36, 48, 60]
    if months not in allowed_months:
        raise HTTPException(status_code=400, detail=f"Invalid months parameter. Allowed values: {allowed_months}")
   
    try:
        df = pd.read_csv(file_path, parse_dates=["Timestamp"])
        filtered_df = df[df["Motor_ID"] == motor_id]
        today = datetime.today()
        start_date = today - relativedelta(months=months)
        filtered_df = filtered_df[(filtered_df["Timestamp"] >= start_date) & (filtered_df["Timestamp"] <= today)]
        if filtered_df.empty:
            return {"message": f"No data available for Motor ID {motor_id} in the past {months} months."}
       
        failure_status = filtered_df[filtered_df["Status"].isin(["High Vibration", "Overload"])]
        trends = failure_status.groupby(failure_status["Timestamp"].dt.to_period("M")).size()
        trends_list = [{"month": str(k), "count": int(v)} for k, v in trends.items()]
       
        if trends_list:
            trends_sorted = sorted(trends_list, key=lambda x: x["month"])
            x_vals = list(range(len(trends_sorted)))
            y_vals = [item["count"] for item in trends_sorted]
            slope, intercept = np.polyfit(x_vals, y_vals, 1)
            regression = {
                "slope": round(slope, 2),
                "intercept": round(intercept, 2),
                "predicted": [
                    {"month": trends_sorted[i]["month"], "predicted_count": round(slope * i + intercept, 2)}
                    for i in range(len(trends_sorted))
                ]
            }
            if len(y_vals) >= 3:
                moving_avg = np.convolve(y_vals, np.ones(3) / 3, mode="valid")
                moving_average = [
                    {"month": trends_sorted[i + 1]["month"], "moving_avg": round(moving_avg[i], 2)}
                    for i in range(len(moving_avg))
                ]
            else:
                moving_average = []
        else:
            regression = {}
            moving_average = []
       
        # Build a prompt for trend analysis of failure data
        failure_data_json = json.dumps(trends_list)
        prompt = (
            f"Analyze the following failure trend data for Motor ID {motor_id}: {failure_data_json}. "
            f"Identify trends, anomalies, and forecast potential future risks. "
             f"Return your result in JSON format with key 'ai_trend'. containing observations and recommendations in bullet points."
        )
        ai_trend_insight = trend_analysis_module.forward(
            analytics_data=failure_data_json, prompt=prompt
        )
       
        return {
            "data": trends_list,
            "total_failures": len(failure_status),
            "period": {
                "start": start_date.strftime("%Y-%m-%d"),
                "end": today.strftime("%Y-%m-%d")
            },
            "ai_trend_insight": ai_trend_insight,
            "trend_analysis": {
                "regression": regression,
                "moving_average": moving_average
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
 
@app.get("/api/motor/rpm-vs-load")
def get_rpm_vs_load(
    motor_id: str = Query(..., description="Motor ID for filtering data"),
    months: int = Query(..., description="Number of months for filtering (1, 6, 12, 24, 36, 48, 60)")
):
    allowed_months = [1, 6, 12, 24, 36, 48, 60]
    if months not in allowed_months:
        raise HTTPException(status_code=400, detail=f"Invalid months parameter. Allowed values: {allowed_months}")
    try:
        df = pd.read_csv(file_path, parse_dates=["Timestamp"])
        for col in ["Motor_ID", "Timestamp", "RPM", "Load (%)"]:
            if col not in df.columns:
                raise HTTPException(status_code=500, detail=f"Missing expected column '{col}' in dataset.")
        filtered_df = df[df["Motor_ID"] == motor_id]
        today = datetime.today()
        start_date = today - relativedelta(months=months)
        filtered_df = filtered_df[(filtered_df["Timestamp"] >= start_date) & (filtered_df["Timestamp"] <= today)]
        if filtered_df.empty:
            return {"message": f"No data available for Motor ID {motor_id} in the past {months} months."}
        filtered_df["Month"] = filtered_df["Timestamp"].dt.to_period("M").astype(str)
        grouped = filtered_df.groupby("Month").agg(
            avg_rpm=("RPM", "mean"),
            avg_load=("Load (%)", "mean")
        ).reset_index()
        grouped["avg_rpm"] = grouped["avg_rpm"].round(2)
        grouped["avg_load"] = grouped["avg_load"].round(2)
       
        # Build aggregated data JSON and prompt for RPM vs. Load trend analysis
        aggregated_data = grouped.to_dict(orient="records")
        aggregated_data_json = json.dumps(aggregated_data)
        prompt = (
            f"Analyze the following monthly aggregated RPM and Load data for Motor ID {motor_id}: {aggregated_data_json}. "
            f"Identify any trends or shifts in performance ."
             f"Return your result in JSON format with key 'ai_trend'. containing observations and recommendations in bullet points."
        )
        ai_trend_insight = trend_analysis_module.forward(
            analytics_data=aggregated_data_json, prompt=prompt
        )
       
        return {
            "data": aggregated_data,
            "period": {
                "start": start_date.strftime("%Y-%m-%d"),
                "end": today.strftime("%Y-%m-%d")
            },
            "ai_trend_insight": ai_trend_insight
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
 
@app.get("/api/motor/temp-vs-vibration")
def get_temp_vs_vibration(
    motor_id: str = Query(..., description="Motor ID for filtering data"),
    months: int = Query(..., description="Number of months for filtering (1, 6, 12, 24, 36, 48, 60)")
):
    allowed_months = [1, 6, 12, 24, 36, 48, 60]
    if months not in allowed_months:
        raise HTTPException(status_code=400, detail=f"Invalid months parameter. Allowed values: {allowed_months}")
    try:
        df = pd.read_csv(file_path, parse_dates=["Timestamp"])
        filtered_df = df[df["Motor_ID"] == motor_id]
        today = datetime.today()
        start_date = today - relativedelta(months=months)
        filtered_df = filtered_df[(filtered_df["Timestamp"] >= start_date) & (filtered_df["Timestamp"] <= today)]
        if filtered_df.empty:
            return {"message": f"No data available for Motor ID {motor_id} in the past {months} months."}
        temp_vib_data = filtered_df[["Temperature (°C)", "Vibration (mm/s)", "Timestamp", "Status"]].copy()
        temp_vib_data["Timestamp"] = temp_vib_data["Timestamp"].dt.strftime("%Y-%m-%d %H:%M:%S")
        correlation = None
        if len(filtered_df) > 1:
            correlation = round(filtered_df[["Temperature (°C)", "Vibration (mm/s)"]].corr().iloc[0, 1], 3)
       
        # Build statistics for trend analysis
        stats = {
            "avg_temperature": round(filtered_df["Temperature (°C)"].mean(), 2),
            "avg_vibration": round(filtered_df["Vibration (mm/s)"].mean(), 2),
            "correlation": correlation,
            "period": {
                "start": start_date.strftime("%Y-%m-%d"),
                "end": today.strftime("%Y-%m-%d")
            }
        }
        stats_json = json.dumps(stats)
        prompt = (
            f"Analyze the following temperature and vibration statistics for Motor ID {motor_id}: {stats_json}. "
            f"Provide an AI-based trend analysis that explains the relationship, highlights any anomalies, and recommends monitoring actions. "
            f"Return your result in JSON format with key 'ai_trend'. containing observations and recommendations in bullet points."
        )
        ai_trend_insight = trend_analysis_module.forward(
            analytics_data=stats_json, prompt=prompt
        )
       
        return {
            "data": temp_vib_data.to_dict(orient="records"),
            "statistics": stats,
            "ai_trend_insight": ai_trend_insight
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# UseCase 2

# Disable warnings and configure logging
urllib3.disable_warnings()
warnings.simplefilter('ignore', InsecureRequestWarning)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Set SSL certificate paths if needed
CERTIFICATE_PATH = os.path.join(os.path.dirname(__file__), "huggingface.co.crt")
os.environ["REQUESTS_CA_BUNDLE"] = CERTIFICATE_PATH
os.environ['CURL_CA_BUNDLE'] = CERTIFICATE_PATH

# Confluence API credentials
CONFLUENCE_BASE_URL = os.getenv("CONFLUENCE_BASE_URL")
CONFLUENCE_USERNAME = os.getenv("CONFLUENCE_USERNAME")
CONFLUENCE_API_TOKEN = os.getenv("CONFLUENCE_API_TOKEN")

# Azure OpenAI credentials (for DSPy and the Azure RAG system)
AZURE_OPENAI_ENDPOINT = os.getenv('AZURE_OPENAI_ENDPOINT')
AZURE_OPENAI_API_KEY = os.getenv('AZURE_OPENAI_API_KEY')
AZURE_OPENAI_API_VERSION = os.getenv('AZURE_OPENAI_API_VERSION')
AZURE_OPENAI_DEPLOYMENT = os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME')

# Azure Blob Storage credentials
AZURE_STORAGE_SAS_URL = os.getenv('AZURE_STORAGE_SAS_URL')

# Vector database path for Azure Blob Storage indexing
VECTOR_DB_PATH = os.getenv("VECTOR_DB_PATH", "vector_db")
UPLOAD_FOLDER = os.getenv("UPLOAD_FOLDER", "vector_db")

# -------------------------------
# DSPy CONFIGURATION
# -------------------------------

# Define DSPy signature for content generation
class GenerateContent(Signature):
    """Generate structured content for a specific section in the specified language."""
    section_title: str = InputField(desc="Title of the section")
    prompt: str = InputField(desc="Prompt for generating content")
    language: str = InputField(desc="Target language for content generation")
    output: str = OutputField(desc="Generated content in specified language")

# -------------------------------
# SSE Progress
# -------------------------------

# Track active tasks for progress updates
active_tasks = {}

# Helper function to update progress
async def update_progress(client_id: str, message: str, percentage: int):
    logger.info(f"Updating progress for {client_id}: {message}, {percentage}%")
    active_tasks[client_id] = {"message": message, "percentage": percentage}
    if percentage >= 100:
        await asyncio.sleep(1)  # Brief delay to ensure client receives final update
        active_tasks.pop(client_id, None)

# -------------------------------
# TRANSLATION & UTILITY FUNCTIONS
# -------------------------------
def load_translations():
    file_path = os.path.join(os.path.dirname(__file__), "translations.json")
    with open(file_path, "r", encoding="utf-8") as f:
        return json.load(f)

TRANSLATIONS = load_translations()

@lru_cache(maxsize=100)
def get_language_texts(language):
    return TRANSLATIONS.get(language, TRANSLATIONS["en"])

def clean_content(text):
    text = re.sub(r'#{1,6}\s*', '', text)
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    text = re.sub(r'\[.*?\]|\{.*?\}', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def clean_product_query(product_name):
    if " - " in product_name:
        product_name = product_name.split(" - ")[0].strip()
    product_name = product_name.replace("®", "")
    product_name = re.sub(r'[^\w\s]', ' ', product_name)
    product_name = re.sub(r'\s+', ' ', product_name)
    return product_name.strip()

# -------------------------------
# CONFLUENCE HANDLING
# -------------------------------
def normalize_text(text):
    """
    Normalize text by removing special characters, extra spaces, and converting to lowercase.
    """
    # Remove special characters and extra spaces
    text = re.sub(r'[^a-zA-Z0-9\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text.lower()

def get_confluence_vector_store(content):
    try:
        if not content.strip():
            return None
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        texts = text_splitter.create_documents([content])
        if not texts:
            return None
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        vector_store = FAISS.from_documents(texts, embeddings)
        return vector_store
    except Exception as e:
        logger.error(f"Error creating Confluence vector store: {str(e)}")
        return None

# -------------------------------
# AZURE BLOB STORAGE INTEGRATION
# -------------------------------
# These classes/functions come from your Azure Blob Storage code with added logging.
class FAISSIndex:
    def __init__(self, dimension: int = 384):
        self.dimension = dimension
        self.index_path = os.path.join(VECTOR_DB_PATH, "index.faiss")
        self.metadata_path = os.path.join(VECTOR_DB_PATH, "metadata.pkl")
        self.documents: List = []
        self.index: Optional[faiss.Index] = None
        os.makedirs(VECTOR_DB_PATH, exist_ok=True)

    def create(self) -> None:
        self.index = faiss.IndexFlatL2(self.dimension)
        logger.info("Created new FAISS index.")

    def add(self, vectors: np.ndarray, documents: List) -> None:
        if self.index is None:
            self.create()
        self.index.add(vectors)
        self.documents.extend(documents)
        logger.info(f"Added {len(documents)} documents to the FAISS index.")

    def save(self) -> None:
        if self.index is not None:
            faiss.write_index(self.index, self.index_path)
            with open(self.metadata_path, 'wb') as f:
                pickle.dump(self.documents, f)
            logger.info("FAISS index and metadata saved successfully.")

    def load(self) -> bool:
        try:
            if os.path.exists(self.index_path) and os.path.exists(self.metadata_path):
                self.index = faiss.read_index(self.index_path)
                with open(self.metadata_path, 'rb') as f:
                    self.documents = pickle.load(f)
                logger.info("FAISS index and metadata loaded successfully.")
                return True
            logger.info("No existing FAISS index found.")
            return False
        except Exception as e:
            logger.error(f"Error loading FAISS index: {str(e)}")
            return False

    def search(self, query_vector: np.ndarray, k: int = 5) -> List[Dict[str, Any]]:
        if self.index is None or not self.documents:
            logger.warning("FAISS index is empty or not loaded.")
            return []
        distances, indices = self.index.search(query_vector.reshape(1, -1), k)
        results = []
        for i, idx in enumerate(indices[0]):
            if idx < 0 or idx >= len(self.documents):
                continue
            doc = self.documents[idx]
            content = doc.page_content
            metadata = doc.metadata
            distance = float(distances[0][i])
            score = 1 / (1 + distance)
            results.append({
                'content': content,
                'metadata': metadata,
                'score': score
            })
        logger.info(f"FAISS search returned {len(results)} results.")
        return results

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    def process_pdf(self, pdf_content: bytes) -> str:
        try:
            # Create a temporary file to handle the PDF content
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(pdf_content)
                temp_file_path = temp_file.name

            # Process the PDF
            with open(temp_file_path, 'rb') as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            
            # Clean up the temporary file
            os.unlink(temp_file_path)
            
            logger.info("Processed PDF content successfully.")
            return text
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            return ""  # Return empty string instead of None

    def create_documents(self, content: str, metadata: Dict[str, Any]) -> List:
        if not content:  # Handle empty content gracefully
            logger.warning("Empty content provided for document creation")
            return []
            
        chunks = self.text_splitter.split_text(content)
        logger.info(f"Split content into {len(chunks)} document chunks.")
        return [Document(page_content=chunk, metadata=metadata) for chunk in chunks]

# A simple Document class (mimicking LangChain's Document)
class Document:
    def __init__(self, page_content: str, metadata: Dict[str, Any]):
        self.page_content = page_content
        self.metadata = metadata

class RAGSystem:
    def __init__(self):
        parsed_url = urllib.parse.urlparse(AZURE_STORAGE_SAS_URL)
        account_name = parsed_url.netloc.split('.')[0]
        container_name = parsed_url.path.strip('/').split('/')[0]
        self.blob_service_client = BlobServiceClient(
            account_url=f"https://{account_name}.blob.core.windows.net",
            credential=AZURE_STORAGE_SAS_URL.split('?')[1],
            connection_verify=False
        )
        self.container_client = self.blob_service_client.get_container_client(container_name)
        logger.info(f"Connected to Azure Blob Storage container: {container_name}")
        self.embeddings = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
        self.index = FAISSIndex()
        self.document_processor = DocumentProcessor()
        from openai import AzureOpenAI
        self.llm = AzureOpenAI(
            api_key=AZURE_OPENAI_API_KEY,
            api_version=AZURE_OPENAI_API_VERSION,
            azure_endpoint=AZURE_OPENAI_ENDPOINT
        )
        self._ensure_index()

    def _ensure_index(self) -> None:
        if not self.index.load():
            logger.info("Building new FAISS index from Azure Blob PDFs.")
            self._build_index()

    def _fetch_blob_content(self) -> List[Dict[str, Any]]:
        documents = []
        logger.info("Fetching blob content from Azure Blob Storage...")
        blob_list = self.container_client.list_blobs()
        for blob in blob_list:
            if blob.name.endswith('.pdf'):
                logger.info(f"Found PDF blob: {blob.name}")
                blob_client = self.container_client.get_blob_client(blob.name)
                blob_content = blob_client.download_blob().readall()
                metadata = {
                    'source': blob.name,
                    'type': 'pdf',
                    'created': blob.creation_time,
                    'modified': blob.last_modified
                }
                documents.append({
                    'content': blob_content,
                    'metadata': metadata
                })
        logger.info(f"Fetched {len(documents)} PDF blobs from Azure.")
        return documents

    def _build_index(self) -> None:
        blobs = self._fetch_blob_content()
        if not blobs:
            logger.warning("No content fetched from Azure Blob Storage")
            return

        all_documents = []
        all_vectors = []

        for blob in blobs:
            try:
                logger.info(f"Processing blob: {blob['metadata'].get('source', 'unknown')}")
                clean_text = self.document_processor.process_pdf(blob['content'])
                documents = self.document_processor.create_documents(clean_text, blob['metadata'])
                vectors = self.embeddings.encode([doc.page_content for doc in documents])
                all_documents.extend(documents)
                all_vectors.extend(vectors)
            except Exception as e:
                logger.error(f"Error processing blob {blob['metadata'].get('source', 'unknown')}: {str(e)}")
                continue

        if all_vectors:
            vectors_np = np.array(all_vectors).astype('float32')
            self.index.add(vectors_np, all_documents)
            self.index.save()
            logger.info("Azure Blob Storage index built successfully.")

    def query(self, query_text: str) -> Dict[str, Any]:
        try:
            logger.info(f"Querying Azure Blob Storage with: {query_text}")
            query_vector = self.embeddings.encode([query_text]).astype('float32')
            search_results = self.index.search(query_vector)
            if not search_results:
                logger.info("No relevant results found in Azure Blob Storage.")
                return {'answer': 'No relevant information found.', 'sources': []}
            context = '\n'.join(str(result['content']) for result in search_results)
            response = self.llm.chat.completions.create(
                model=AZURE_OPENAI_DEPLOYMENT,
                messages=[
                    {"role": "system", "content": "You are a presales expert. Provide accurate, concise answers based only on the provided context from Azure Blob Storage PDFs. Do not use any external knowledge."},
                    {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {query_text}"}
                ],
                temperature=0.7,
                max_tokens=500
            )
            seen_sources = set()
            sources = []
            for result in search_results:
                source_title = str(result['metadata'].get('source', 'Unknown'))
                if source_title not in seen_sources:
                    source = {
                        'title': source_title,
                        'confidence': round(float(result['score']) * 100, 2),
                        'modified': result['metadata'].get('modified', 'Unknown')
                    }
                    sources.append(source)
                    seen_sources.add(source_title)
            logger.info("Azure Blob Storage query processed successfully.")
            return {
                'answer': response.choices[0].message.content.strip(),
                'sources': sources
            }
        except Exception as e:
            logger.error(f"Error in Azure Blob query processing: {str(e)}")
            return {
                'answer': 'An error occurred while processing your query.',
                'sources': [],
                'error': str(e)
            }

def get_blob_client(blob_name: str):
    """Helper function to get blob client"""
    parsed_url = urllib.parse.urlparse(AZURE_STORAGE_SAS_URL)
    account_name = parsed_url.netloc.split('.')[0]
    container_name = parsed_url.path.strip('/').split('/')[0]
    blob_service_client = BlobServiceClient(
        account_url=f"https://{account_name}.blob.core.windows.net",
        credential=AZURE_STORAGE_SAS_URL.split('?')[1],
        connection_verify=False
    )
    container_client = blob_service_client.get_container_client(container_name)
    return container_client.get_blob_client(blob_name)

async def convert_to_pdf(file: UploadFile) -> bytes:
    """Convert uploaded file to PDF format"""
    try:
        content = await file.read()
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension == '.pdf':
            logger.info(f"File {file.filename} is already in PDF format")
            return content
        
        elif file_extension in ['.docx', '.doc']:
            logger.info(f"Starting conversion of {file_extension} file: {file.filename} to PDF")
            
            # Create a temporary directory
            temp_dir = tempfile.mkdtemp()
            temp_doc_path = os.path.join(temp_dir, f"document{file_extension}")
            temp_pdf_path = os.path.join(temp_dir, "document.pdf")
            
            # Write the document file to disk
            with open(temp_doc_path, 'wb') as f:
                f.write(content)
                
            logger.info(f"Created temporary file at {temp_doc_path}")
            
            # Initialize COM in a separate thread
            pythoncom.CoInitialize()
            success = False
            
            try:
                logger.info("Creating Word application instance")
                word = win32com.client.DispatchEx('Word.Application')
                word.Visible = False
                word.DisplayAlerts = 0
                
                try:
                    # Use absolute paths
                    abs_doc_path = os.path.abspath(temp_doc_path)
                    abs_pdf_path = os.path.abspath(temp_pdf_path)
                    
                    logger.info(f"Opening document from {abs_doc_path}")
                    doc = word.Documents.Open(abs_doc_path, ReadOnly=1)
                    
                    logger.info(f"Saving as PDF to {abs_pdf_path}")
                    doc.SaveAs(abs_pdf_path, FileFormat=17)  # 17 = wdFormatPDF
                    doc.Close()
                    
                    if os.path.exists(abs_pdf_path):
                        logger.info("PDF created successfully")
                        with open(abs_pdf_path, 'rb') as pdf_file:
                            success = True
                            return pdf_file.read()
                    else:
                        logger.error(f"PDF file not found at {abs_pdf_path}")
                        
                except Exception as e:
                    logger.error(f"Error in Word automation: {str(e)}", exc_info=True)
                    if "RPC_E_SERVERCALL_RETRYLATER" in str(e):
                        logger.error("RPC server busy - Word might be running in non-interactive mode")
                    elif "Call was rejected by callee" in str(e):
                        logger.error("COM call rejected - could be security settings or privileges")
                
                finally:
                    try:
                        word.Quit()
                    except:
                        pass
            
            except Exception as e:
                logger.error(f"Error creating Word application: {str(e)}", exc_info=True)
            
            finally:
                pythoncom.CoUninitialize()
                
                # Clean up temp files
                try:
                    import shutil
                    shutil.rmtree(temp_dir)
                    logger.info(f"Cleaned up temporary directory {temp_dir}")
                except Exception as e:
                    logger.warning(f"Failed to clean up: {str(e)}")
            
            if not success:
                raise Exception(f"Failed to convert {file_extension} to PDF using COM automation")
                
        elif file_extension == '.txt':
            logger.info(f"Starting conversion of TXT file: {file.filename} to PDF")
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            text_content = content.decode('utf-8')
            lines = text_content.split('\n')
            
            for line in lines:
                # Replace any non-printable characters
                line = ''.join(c if ord(c) < 128 else ' ' for c in line)
                if line.strip():
                    pdf.multi_cell(0, 10, txt=line)
                else:
                    pdf.ln(5)
            
            pdf_content = pdf.output(dest='S').encode('latin-1')
            logger.info(f"Successfully converted TXT file: {file.filename} to PDF")
            return pdf_content
            
        else:
            logger.error(f"Unsupported file type: {file_extension}")
            raise HTTPException(
                status_code=400,
                detail="Supported file types: PDF, DOCX, DOC, and TXT files."
            )
    
    except Exception as e:
        logger.error(f"Error converting file to PDF: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to convert file to PDF: {str(e)}"
        )

async def upload_to_azure_blob(file: UploadFile) -> str:
    """Upload file to Azure Blob Storage"""
    try:
        pdf_content = await convert_to_pdf(file)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        original_name = os.path.splitext(secure_filename(file.filename))[0]
        blob_name = f"{timestamp}_{original_name}.pdf"
        
        blob_client = get_blob_client(blob_name)
        blob_client.upload_blob(pdf_content, overwrite=True)
        
        logger.info(f"Successfully uploaded converted PDF file {blob_name} to Azure Blob Storage")
        return blob_name
        
    except Exception as e:
        logger.error(f"Error uploading to Azure Blob Storage: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to upload file: {str(e)}")

# Helper function to retrieve Azure Blob Storage content for a given query
def retrieve_azure_blob_content(query: str) -> str:
    try:
        logger.info(f"Retrieving Azure Blob content for query: {query}")
        azure_system = RAGSystem()
        result = azure_system.query(query)
        content = result.get('answer', '')
        if not content.strip():
            return "No relevant Azure Blob Storage content found."
        return content
    except Exception as e:
        logger.error(f"Error retrieving Azure blob content: {str(e)}")
        return "No relevant Azure Blob Storage content found."

# -------------------------------
# COMBINING ALL SOURCES
# -------------------------------
def combine_all_content(scraped_data, pdf_content, confluence_content, azure_blob_content):
    combined_content = []
    used_content = set()  # Track used content to avoid duplicates

    def add_content(section_title, content):
        if content and isinstance(content, str) and content.strip() and content not in used_content:
            combined_content.append(f"=== {section_title} ===")
            combined_content.append(content)
            used_content.add(content)

    # Add scraped data
    if scraped_data:
        product_info = []
        product_info.append(f"Product Name: {scraped_data.get('product_name', 'N/A')}")
        if features := scraped_data.get('key_features', []):
            product_info.append("\nKey Features:")
            product_info.extend([f"- {feature}" for feature in features])
        add_content("Product Information", "\n".join(product_info))

        # Add specifications separately
        if tech_specs := scraped_data.get('technical_specifications', {}):
            specs_content = []
            specs_content.append("Technical Specifications:")
            for key, value in tech_specs.items():
                specs_content.append(f"- {key}: {value}")
            add_content("Technical Specifications", "\n".join(specs_content))

        if gen_specs := scraped_data.get('general_specifications', {}):
            specs_content = []
            specs_content.append("General Specifications:")
            for key, value in gen_specs.items():
                specs_content.append(f"- {key}: {value}")
            add_content("General Specifications", "\n".join(specs_content))

    # Add PDF content if it exists and is a string
    if pdf_content and isinstance(pdf_content, str):
        add_content("Additional Documentation", pdf_content)

    # Add Confluence content if it exists and is a string
    if confluence_content and isinstance(confluence_content, str):
        add_content("Confluence Documentation", confluence_content)

    # Add Azure Blob content if it exists and is a string
    if azure_blob_content and isinstance(azure_blob_content, str):
        add_content("Azure Documentation", azure_blob_content)

    # If no content was added, add a default message
    if not combined_content:
        combined_content.append("No content available for this product.")

    return "\n\n".join(combined_content)

# -------------------------------
# PDF GENERATION & WEB SCRAPING
# -------------------------------
def generate_pdf(product_data, content, is_faq=False):
    try:
        buffer = BytesIO()
        
        # Dictionary to store section names and their page numbers
        section_pages = {}
        
        # Custom paragraph class to set bookmarks
        class BookmarkParagraph(Paragraph):
            def __init__(self, text, style, bookmark_key):
                super().__init__(text, style)
                self.bookmark_key = bookmark_key
            
            def draw(self):
                super().draw()
                self.canv.bookmarkPage(self.bookmark_key)
                section_pages[self.bookmark_key] = self.canv.getPageNumber()
        
        # Function to add page numbers to ALL pages
        def add_page_number(canvas, doc):
            canvas.saveState()
            # Draw the black border
            canvas.setStrokeColor(colors.black)  # Set border color to black
            canvas.setLineWidth(1)               # Set line thickness to 1 point
            border_margin = 28
            canvas.rect(
                border_margin,                   # x-coordinate (left edge)
                border_margin,                   # y-coordinate (bottom edge)
                doc.pagesize[0] - 2 * border_margin,  # Width of the rectangle
                doc.pagesize[1] - 2 * border_margin   # Height of the rectangle
            )
            page_num = canvas.getPageNumber()
            text = f"Page {page_num}"
            canvas.setFont('Helvetica', 10)
            canvas.drawCentredString(doc.pagesize[0] / 2, 36, text)
            logger.info(f"Drawing page number {page_num} at y=36")
            canvas.restoreState()
        
        # Build document with helper function
        def create_doc():
            # Create the document
            doc = BaseDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Create a single frame for all pages
            frame = Frame(
                doc.leftMargin,
                doc.bottomMargin + 40,  # Leave space for page numbers
                doc.width,
                doc.height - 40,
                id='normal'
            )
            
            # Create a SINGLE template for ALL pages
            template = PageTemplate(
                id='all_pages', 
                frames=frame,
                onPage=add_page_number
            )
            
            # Add ONLY this template to the document - no defaults
            doc.addPageTemplates([template])
            
            return doc
        
        # Get styles
        styles = getSampleStyleSheet()
        
        title_style = styles['Title']
        title_style.fontName = 'Helvetica-Bold'
        title_style.fontSize = 18
        title_style.textColor = colors.HexColor('#1e40af')
        
        heading1_style = styles['Heading1']
        heading1_style.fontName = 'Helvetica-Bold'
        heading1_style.fontSize = 16
        heading1_style.textColor = colors.HexColor('#1e3a8a')
        
        heading2_style = styles['Heading2']
        heading2_style.fontName = 'Helvetica-Bold'
        heading2_style.fontSize = 14
        heading2_style.textColor = colors.HexColor('#2563eb') if not is_faq else colors.black
        
        normal_style = styles['Normal']
        normal_style.fontName = 'Helvetica'
        normal_style.fontSize = 11
        normal_style.leading = 14
        normal_style.textColor = colors.black
        
        # Build elements list
        def build_elements(include_toc_pages=True):
            elements = []
            
            language_texts = get_language_texts(product_data.get("language", "en"))
            if is_faq:
                title_text = f"{language_texts['faq_title']}"
            else:
                title_text = f"{language_texts['manual_title']}"
            
            elements.append(Paragraph(title_text, title_style))
            elements.append(Spacer(1, 0.25 * inch))
            elements.append(Paragraph(product_data['product_category'], styles['Heading3']))
            elements.append(Spacer(1, 0.5 * inch))
            
            elements.append(Paragraph(language_texts['table_of_contents'], heading1_style))
            toc_data = [[language_texts['section'], language_texts['page']]]
            for section in content.keys():
                clean_section = clean_content(section)
                page_num = section_pages.get(clean_section, "") if include_toc_pages else ""
                toc_data.append([clean_section, str(page_num)])
            
            toc_table = Table(toc_data, colWidths=[400, 100])
            toc_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 13),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 15),
                ('GRID', (0, 0), (-1, -1), 1, colors.lightgrey),
                ('BOX', (0, 0), (-1, -1), 2, colors.HexColor('#cbd5e1')),
                ('LEFTPADDING', (0, 0), (-1, -1), 15),
                ('RIGHTPADDING', (0, 0), (-1, -1), 15),
            ]))
            elements.append(toc_table)
            
            for section, section_content in content.items():
                elements.append(PageBreak())
                clean_section = clean_content(section)
                elements.append(BookmarkParagraph(clean_section, heading1_style, clean_section))
                elements.append(Spacer(1, 0.1 * inch))

                if section == language_texts["technical_specifications"]:
                    tables = format_specifications_tables(product_data, is_faq)
                    if tables:
                        for table in tables:
                            elements.append(table)
                            elements.append(Spacer(1, 0.2 * inch))
                        continue
                
                paragraphs = clean_content(section_content).split('\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        if paragraph.strip().endswith(':'):
                            elements.append(Paragraph(paragraph.strip(), heading2_style))
                        else:
                            elements.append(Paragraph(paragraph.strip(), normal_style))
                        elements.append(Spacer(1, 0.05 * inch))
                
                elements.append(Spacer(1, 0.2 * inch))
            
            return elements
        
        # First build to collect page numbers
        doc = create_doc()
        first_elements = build_elements(include_toc_pages=False)
        doc.build(first_elements)
        
        # Second build with updated TOC
        buffer.seek(0)
        buffer.truncate(0)  # Clear the buffer for the second build
        doc = create_doc()
        second_elements = build_elements(include_toc_pages=True)
        doc.build(second_elements)
        
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")

def get_product_link(selected_item):
    for product in products_data.get("products", []):
        for subproduct in product.get("subproducts", []):
            for item in subproduct.get("sub_subproducts", []):
                if item.get("sub_subproduct_name") == selected_item:
                    return item.get("sub_subproduct_link")
    return None

def generate_content_prompts(cleaned_product_name, combined_content, language):
    language_texts = get_language_texts(language)
    language_instruction = (
        f"You are a professional technical writer creating content in {language}.\n"
        "Instructions:\n"
        "1. Generate ALL content in the target language.\n"
        "2. Maintain technical accuracy and use a formal tone.\n"
        "3. Preserve all technical terms and measurements.\n"
        "4. Keep the same structured format as the original.\n"
        "5. Ensure all headings and subheadings are in the target language.\n"
        "6. IMPORTANT: Create unique content for each section that doesn't duplicate information from other sections.\n"
    )
    context_text = f"\n\nRelevant context:\n{combined_content}\n\n"
    prompts = {}
    sections = {
        "introduction": language_texts["introduction"],
        "key_features": language_texts["key_features"],
        "technical_specifications": language_texts["technical_specifications"],
        "safety_information": language_texts["safety_information"],
        "setup_instructions": language_texts["setup_instructions"],
        "operation_instructions": language_texts["operation_instructions"],
        "maintenance_and_care": language_texts["maintenance_and_care"],
        "troubleshooting": language_texts["troubleshooting"],
        "warranty_information": language_texts["warranty_information"]
    }
    
    # Special instructions for sections that often overlap
    section_specific_instructions = {
        "maintenance_and_care": "IMPORTANT: Focus on regular maintenance tasks like cleaning, lubrication, and inspection. Do not include content about fixing problems or diagnosing issues, as that belongs in Troubleshooting.",
        "troubleshooting": "IMPORTANT: Focus on diagnosing and fixing specific problems or issues. Do not include routine maintenance tasks, as those belong in Maintenance and Care.",
        "technical_specifications": "IMPORTANT: This section should consist of tabular data and precise measurements. Do not repeat detailed descriptions of features.",
        "key_features": "IMPORTANT: Focus on the most important capabilities and benefits. Do not include detailed specifications as those belong in Technical Specifications."
    }
    
    for key, section_title in sections.items():
        prompt = f"{language_instruction}{context_text}"
        
        # Add section-specific instructions to reduce redundancy
        if key in section_specific_instructions:
            prompt += f"{section_specific_instructions[key]}\n\n"
            
        prompt += f"Task: Generate a detailed '{section_title}' section for {cleaned_product_name} in {language}."
        prompts[section_title] = prompt
        
    return prompts

async def translate_specifications(specs: Dict[str, str], language: str) -> Dict[str, str]:
    """Translate specification keys and values into the target language using DSPy."""
    try:
        if not specs:
            return {}
        
        # Prepare the prompt for translation
        specs_text = "\n".join([f"{key}: {value}" for key, value in specs.items()])
        language_texts = get_language_texts(language)
        prompt = f"""
        You are a professional translator converting technical specifications into {language}.
        Instructions:
        1. Translate the following specification keys and values into {language}.
        2. Preserve technical accuracy and maintain a formal tone.
        3. Do not translate units (e.g., 'V', 'rpm', 'Hz', 'LB', 'IN') or proper nouns (e.g., brand names, country names like 'Mexico').
        4. Return the translated content in the same key-value format.
        
        Specifications to translate:
        {specs_text}
        """
        
        # Use DSPy to translate
        predictor = Predict(GenerateContent)
        result = await asyncio.to_thread(
            lambda: predictor(
                section_title=f"Translated Specifications in {language}",
                prompt=prompt,
                language=language
            )
        )
        
        if not result or not hasattr(result, 'output'):
            logger.warning(f"Failed to translate specifications into {language}")
            return specs  # Fallback to original if translation fails
        
        # Parse the translated output back into a dictionary
        translated_specs = {}
        lines = result.output.strip().split('\n')
        for line in lines:
            if ':' in line:
                key, value = line.split(':', 1)
                translated_specs[key.strip()] = value.strip()
        
        logger.info(f"Translated {len(translated_specs)} specification items into {language}")
        return translated_specs
    
    except Exception as e:
        logger.error(f"Error translating specifications into {language}: {str(e)}")
        return specs  # Fallback to original on error

def format_specifications_tables(product_data, is_faq=False):
    try:
        tables = []
        styles = getSampleStyleSheet()
        language = product_data.get("language", "en")
        language_texts = get_language_texts(language)
        
        sub_header_style = styles['Heading4']
        sub_header_style.fontName = 'Helvetica-Bold'
        sub_header_style.fontSize = 12
        sub_header_style.textColor = colors.HexColor('#1e40af') if not is_faq else colors.black
        
        header_bg_color = colors.HexColor('#e6efff') if not is_faq else colors.HexColor('#f5f5f5')
        header_text_color = colors.HexColor('#1e40af') if not is_faq else colors.black
        
        scraped_data = product_data.get("scraped_data", {})
        
        tech_specs = scraped_data.get('technical_specifications', {})
        if language != "en":
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                tech_specs = loop.run_until_complete(translate_specifications(tech_specs, language))
            finally:
                loop.close()
        
        if tech_specs:
            logger.info(f"Formatting {len(tech_specs)} technical specifications")
            tables.append(Paragraph(language_texts["technical_specifications"], sub_header_style))
            tables.append(Spacer(1, 0.1 * inch))
            
            # Create paragraph style for cell content with wrapping
            cell_style = styles['Normal'].clone('CellStyle')
            cell_style.fontSize = 10
            cell_style.leading = 12  # Line spacing
            
            # Prepare data with paragraphs to enable wrapping
            data = [[Paragraph(language_texts["specification"], cell_style), 
                     Paragraph(language_texts["value"], cell_style)]]
            
            for key, value in tech_specs.items():
                data.append([
                    Paragraph(str(key), cell_style), 
                    Paragraph(str(value) if value is not None else "N/A", cell_style)
                ])
            
            if len(data) > 1:
                # Adjust column widths (first column wider for specification names)
                table = Table(data, colWidths=[275, 225])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), header_bg_color),
                    ('TEXTCOLOR', (0, 0), (-1, 0), header_text_color),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
                    ('BOX', (0, 0), (-1, -1), 2, colors.HexColor('#cbd5e1')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Vertical alignment
                    ('TOPPADDING', (0, 1), (-1, -1), 8),     # Add more padding between rows
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                ]))
                tables.append(table)
            else:
                logger.warning("No valid technical specifications data to format")
        
        gen_specs = scraped_data.get('general_specifications', {})
        if language != "en":
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                gen_specs = loop.run_until_complete(translate_specifications(gen_specs, language))
            finally:
                loop.close()
        
        if gen_specs:
            logger.info(f"Formatting {len(gen_specs)} general specifications")
            tables.append(Spacer(1, 0.5 * inch))
            tables.append(Paragraph(language_texts["general_specifications"], sub_header_style))
            tables.append(Spacer(1, 0.1 * inch))
            
            # Create paragraph style for cell content with wrapping
            cell_style = styles['Normal'].clone('CellStyle')
            cell_style.fontSize = 10
            cell_style.leading = 12  # Line spacing
            
            # Prepare data with paragraphs to enable wrapping
            data = [[Paragraph(language_texts["specification"], cell_style), 
                     Paragraph(language_texts["value"], cell_style)]]
            
            for key, value in gen_specs.items():
                data.append([
                    Paragraph(str(key), cell_style), 
                    Paragraph(str(value) if value is not None else "N/A", cell_style)
                ])
            
            if len(data) > 1:
                # Adjust column widths (first column wider for specification names)
                table = Table(data, colWidths=[275, 225])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), header_bg_color),
                    ('TEXTCOLOR', (0, 0), (-1, 0), header_text_color),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
                    ('BOX', (0, 0), (-1, -1), 2, colors.HexColor('#cbd5e1')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Vertical alignment
                    ('TOPPADDING', (0, 1), (-1, -1), 8),     # Add more padding between rows
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                ]))
                tables.append(table)
            else:
                logger.warning("No valid general specifications data to format")
        
        if tables:
            logger.info(f"Formatted {len(tables)} tables for specifications")
            return tables
        else:
            logger.info("No specification tables created")
            return None
    except Exception as e:
        logger.error(f"Error formatting specification tables: {str(e)}")
        return None
    
def run_generate_content(section_title: str, prompt: str, language: str) -> Tuple[str, str]:
    """Generate content for a specific section."""
    try:
        generate_content = Predict(GenerateContent)
        result = generate_content(
            section_title=section_title,
            prompt=prompt,
            language=language
        )
        return section_title, result.output
    except Exception as e:
        logger.error(f"Error generating content for {section_title}: {str(e)}")
        return section_title, ""

async def parallel_content_generation(prompts: Dict[str, str], language: str) -> Dict[str, str]:
    try:
        # Create a ThreadPoolExecutor instead of ProcessPoolExecutor
        with concurrent.futures.ThreadPoolExecutor() as executor:
            loop = asyncio.get_event_loop()
            
            # Create tasks for each prompt
            futures = []
            for title, prompt in prompts.items():
                future = loop.run_in_executor(
                    executor,
                    run_generate_content,
                    title,
                    prompt,
                    language
                )
                futures.append(future)
            
            # Wait for all tasks to complete
            completed_results = await asyncio.gather(*futures, return_exceptions=True)
            
            # Process results
            result_dict = {}
            for result in completed_results:
                if isinstance(result, Exception):
                    logger.error(f"Error in content generation: {str(result)}")
                    continue
                if isinstance(result, tuple) and len(result) == 2:
                    section_title, content = result
                    if content and content.strip():  # Only add non-empty content
                        result_dict[section_title] = content
            
            if not result_dict:
                logger.error("No content was generated successfully")
                raise ValueError("Failed to generate any content")
                
            return result_dict

    except Exception as e:
        logger.error(f"Error in parallel content generation: {str(e)}")
        raise ValueError(f"Content generation failed: {str(e)}")

@app.post("/api/motor/generate-manual")
async def generate_manual(
    product_category: str = Form(...),
    rag_source: Optional[UploadFile] = File(None),
    language: str = Form(...),
    client_id: str = Form(...)
):
    try:
        logger.info(f"Starting generation for client_id: {client_id}")
        await update_progress(client_id, "Initializing document generation...", 5)

        async with aiohttp.ClientSession() as session:
            await update_progress(client_id, "Retrieving product information...", 10)
            product_link = get_product_link(product_category)
            if not product_link:
                raise HTTPException(status_code=400, detail="Product link not found")

            await update_progress(client_id, "Gathering information from sources...", 20)
            tasks = [
                async_scrape_product_data(product_link, session),
                async_search_confluence(product_category, session)
            ]
            scraped_data, confluence_content = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Handle potential exceptions in tasks
            if isinstance(scraped_data, Exception):
                logger.error(f"Scraping failed: {str(scraped_data)}")
                scraped_data = {
                    "product_name": "Unknown Product",
                    "key_features": [],
                    "technical_specifications": {},
                    "general_specifications": {}
                }
            if isinstance(confluence_content, Exception):
                logger.error(f"Confluence search failed: {str(confluence_content)}")
                confluence_content = ""

            cleaned_product_name = clean_product_query(scraped_data["product_name"])

            if rag_source:
                await update_progress(client_id, "Processing uploaded file...", 30)
                try:
                    await upload_to_azure_blob(rag_source)
                except Exception as e:
                    logger.error(f"Failed to upload PDF: {str(e)}")

            await update_progress(client_id, "Retrieving additional content...", 40)
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(retrieve_azure_blob_content, cleaned_product_name)
                azure_blob_content = await asyncio.get_event_loop().run_in_executor(None, lambda: future.result())

            await update_progress(client_id, "Analyzing content...", 50)
            combined_content = combine_all_content(scraped_data, "", confluence_content, azure_blob_content)

            await update_progress(client_id, "Preparing content generation...", 60)
            prompts = generate_content_prompts(cleaned_product_name, combined_content, language)

            try:
                await update_progress(client_id, "Generating manual content...", 70)
                generated_content = await parallel_content_generation(prompts, language)
                if not generated_content:
                    raise HTTPException(status_code=500, detail="No content was generated successfully")
            except ValueError as ve:
                raise HTTPException(status_code=500, detail=str(ve))
            except Exception as e:
                logger.error(f"Unexpected error in content generation: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Content generation failed: {str(e)}")

            try:
                await update_progress(client_id, "Creating PDF document...", 85)
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    pdf_buffer = await asyncio.get_event_loop().run_in_executor(
                        executor,
                        generate_pdf,
                        {"product_category": product_category, "product_name": scraped_data["product_name"], "language": language, "scraped_data": scraped_data},
                        generated_content
                    )
            except Exception as e:
                logger.error(f"Error generating PDF: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")

            await update_progress(client_id, "Finalizing document...", 95)
            filename = f"user_manual_{scraped_data['product_name']}_{language}.pdf"
            encoded_filename = quote(filename)
            response = StreamingResponse(pdf_buffer, media_type="application/pdf")
            response.headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{encoded_filename}"

            await update_progress(client_id, "Document ready!", 100)
            return response

    except HTTPException as he:
        logger.error(f"HTTP exception in manual generation: {str(he)}")
        await update_progress(client_id, "Error occurred", 100)
        raise he
    except Exception as e:
        logger.error(f"Error in manual generation: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        await update_progress(client_id, "Error occurred", 100)
        raise HTTPException(status_code=500, detail=f"Manual generation failed: {str(e)}")
    finally:
        active_tasks.pop(client_id, None)

@app.post("/api/motor/generate-faq")
async def generate_faq(
    product_category: str = Form(...),
    language: str = Form(...),
    preview: bool = Form(True),
    client_id: str = Form(...)
):
    try:
        logger.info(f"Starting FAQ generation for client_id: {client_id}")
        await update_progress(client_id, "Initializing FAQ generation...", 5)

        async with aiohttp.ClientSession() as session:
            await update_progress(client_id, "Retrieving product information...", 10)
            product_link = get_product_link(product_category)
            if not product_link:
                raise HTTPException(status_code=400, detail="Product link not found")

            await update_progress(client_id, "Gathering information from sources...", 20)
            tasks = [
                async_scrape_product_data(product_link, session),
                async_search_confluence(product_category, session)
            ]
            scraped_data, confluence_content = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Handle potential exceptions in tasks
            if isinstance(scraped_data, Exception):
                logger.error(f"Scraping failed: {str(scraped_data)}")
                scraped_data = {
                    "product_name": "Unknown Product",
                    "key_features": [],
                    "technical_specifications": {},
                    "general_specifications": {}
                }
            if isinstance(confluence_content, Exception):
                logger.error(f"Confluence search failed: {str(confluence_content)}")
                confluence_content = ""

            cleaned_product_name = clean_product_query(scraped_data["product_name"])

            await update_progress(client_id, "Retrieving additional content...", 30)
            azure_blob_content = await asyncio.get_event_loop().run_in_executor(
                None, retrieve_azure_blob_content, cleaned_product_name
            )

            language_texts = get_language_texts(language)

            try:
                await update_progress(client_id, "Analyzing data and preparing FAQ content...", 40)
                predictor = Predict(GenerateContent)
                input_data = {
                    "section_title": language_texts["faq"],
                    "prompt": f"""Generate a comprehensive FAQ section for {cleaned_product_name}.
                    Include questions and answers about:
                    - Product features and specifications
                    - Installation and setup
                    - Common usage scenarios
                    - Troubleshooting
                    - Maintenance and care
                    
                    Product Information:
                    {json.dumps(scraped_data, indent=2)}
                    
                    Additional Context:
                    {confluence_content}
                    
                    Azure Blob Storage Content:
                    {azure_blob_content}
                    """,
                    "language": language
                }
                await update_progress(client_id, "Generating FAQ content...", 60)
                result = await asyncio.to_thread(lambda: predictor(**input_data))
                if not result or not hasattr(result, 'output'):
                    logger.error("No FAQ content was generated")
                    raise HTTPException(status_code=500, detail="No FAQ content was generated")
                logger.info("FAQ content generated successfully")
            except Exception as e:
                logger.error(f"Error generating FAQ content: {str(e)}")
                raise HTTPException(status_code=500, detail=f"FAQ generation failed: {str(e)}")

            try:
                await update_progress(client_id, "Creating PDF document...", 80)
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    pdf_buffer = await asyncio.get_event_loop().run_in_executor(
                        executor,
                        generate_pdf,
                        {"product_category": product_category, "product_name": scraped_data["product_name"], "language": language, "scraped_data": scraped_data},
                        {language_texts["faq"]: result.output},
                        True
                    )
                logger.info("PDF generated successfully")
                await update_progress(client_id, "Finalizing document...", 95)
            except Exception as e:
                logger.error(f"Error generating PDF: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")

            if preview:
                import base64
                pdf_bytes = pdf_buffer.getvalue()
                pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
                await update_progress(client_id, "Document ready!", 100)
                return JSONResponse({"pdf_base64": pdf_base64, "filename": f"faq_{scraped_data['product_name']}_{language}.pdf"})
            else:
                filename = f"faq_{scraped_data['product_name']}_{language}.pdf"
                response = StreamingResponse(pdf_buffer, media_type="application/pdf")
                response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
                await update_progress(client_id, "Document ready!", 100)
                return response

    except HTTPException as he:
        logger.error(f"HTTP exception in FAQ generation: {str(he)}")
        await update_progress(client_id, "Error occurred", 100)
        raise he
    except Exception as e:
        logger.error(f"Error in FAQ generation: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        await update_progress(client_id, "Error occurred", 100)
        raise HTTPException(status_code=500, detail=f"FAQ generation failed: {str(e)}")
    finally:
        active_tasks.pop(client_id, None)

PRODUCTS_FILE_PATH = os.path.join(os.path.dirname(__file__), "product_names.json")
with open(PRODUCTS_FILE_PATH, "r") as file:
    products_data = json.load(file)

@app.get("/api/motor/products")
async def get_products():
    return JSONResponse(content={"products": products_data.get("products", [])})

async def async_scrape_product_data(url: str, session: aiohttp.ClientSession) -> Dict[str, Any]:
    """Async version of scrape_product_data with dynamic headers and robust error handling"""
    try:
        # Initialize UserAgent for dynamic rotation
        ua = UserAgent()
        headers = {
            "User-Agent": ua.random,  # Start with a random User-Agent
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Accept-Encoding": "gzip, deflate, br",
            "Connection": "keep-alive",
            "Upgrade-Insecure-Requests": "1",
            "Referer": "https://www.google.com/",
        }
        
        max_retries = 3
        content = None
        for attempt in range(max_retries):
            try:
                timeout = aiohttp.ClientTimeout(total=30)
                # Rotate User-Agent for each retry
                headers["User-Agent"] = ua.random
                logger.info(f"Attempt {attempt + 1}/{max_retries} to scrape {url} with User-Agent: {headers['User-Agent']}")
                async with session.get(url, headers=headers, verify_ssl=False, timeout=timeout) as response:
                    if response.status == 403 or response.status == 429:
                        retry_after = int(response.headers.get('Retry-After', 2 ** attempt))
                        logger.warning(f"Received {response.status} for {url}, retrying after {retry_after}s (attempt {attempt + 1}/{max_retries})")
                        await asyncio.sleep(retry_after)
                        continue
                    response.raise_for_status()
                    content = await response.text()
                    logger.info(f"Successfully retrieved content from {url}")
                    break
            except aiohttp.ClientResponseError as e:
                if attempt == max_retries - 1:
                    logger.error(f"Scraping failed after {max_retries} attempts: {str(e)}")
                    return {
                        "product_name": "Unknown Product",
                        "key_features": [],
                        "technical_specifications": {},
                        "general_specifications": {}
                    }
            except Exception as e:
                logger.error(f"Unexpected error on attempt {attempt + 1}: {str(e)}")
                if attempt == max_retries - 1:
                    return {
                        "product_name": "Unknown Product",
                        "key_features": [],
                        "technical_specifications": {},
                        "general_specifications": {}
                    }
        
        # If no content was retrieved after retries, return fallback
        if content is None:
            logger.error(f"No content retrieved from {url} after {max_retries} attempts")
            return {
                "product_name": "Unknown Product",
                "key_features": [],
                "technical_specifications": {},
                "general_specifications": {}
            }
        
        # Parse the content with BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')
        
        # Extract product name
        product_name = "Unknown Product"
        h1_tag = soup.find('h1')
        if h1_tag:
            product_name = h1_tag.get_text(strip=True)
        logger.info(f"Scraped product name: {product_name}")
        
        # Extract key features
        key_features = []
        key_features_container = soup.find('div', class_='product-info')
        if key_features_container:
            feature_list = key_features_container.find('ul')
            if feature_list:
                features = feature_list.find_all('li')
                for feature in features:
                    key_features.append(feature.get_text(strip=True))
        logger.info(f"Scraped {len(key_features)} key features")
        
        # Initialize dictionaries for specifications
        technical_specs = {}
        general_specs = {}
        
        # Find the specification navigation links
        spec_nav = soup.find('ul', class_='pdp-spec-nav')
        if not spec_nav:
            logger.warning("No pdp-spec-nav found; unable to categorize specifications")
            return {
                "product_name": product_name,
                "key_features": key_features,
                "technical_specifications": technical_specs,
                "general_specifications": general_specs
            }
        
        # Map tab labels to their IDs
        tab_mapping = {}
        for nav_item in spec_nav.find_all('a', class_='pdp-spec-nav__item'):
            tab_label = nav_item.get_text(strip=True).lower()
            tab_id = nav_item.get('href', '').lstrip('#')  # e.g., "tab-0"
            if tab_id:
                tab_mapping[tab_id] = tab_label
                logger.info(f"Found tab mapping: {tab_id} -> {tab_label}")
        
        # Find the tab content container
        tab_content = soup.find('div', class_='tab-content')
        if not tab_content:
            logger.warning("No tab-content div found; cannot process specifications")
            return {
                "product_name": product_name,
                "key_features": key_features,
                "technical_specifications": technical_specs,
                "general_specifications": general_specs
            }
        logger.info(f"Found tab-content div")
        
        # Process each tab by ID from tab_mapping
        tabs = [soup.find('div', id=tab_id) for tab_id in tab_mapping.keys()]
        tabs = [tab for tab in tabs if tab is not None]  # Filter out None results
        logger.info(f"Found {len(tabs)} tabs with matching IDs: {[tab.get('id') for tab in tabs]}")
        
        for tab in tabs:
            tab_id = tab.get('id')
            if tab_id not in tab_mapping:
                logger.warning(f"Tab {tab_id} has no corresponding nav link; skipping")
                continue
                
            tab_label = tab_mapping[tab_id]
            specs_table = tab.find('table', class_='specifications-table')
            if not specs_table:
                logger.warning(f"No specifications table found in {tab_id} ({tab_label})")
                continue
                
            # Extract specifications from the table
            specs_dict = {}
            rows = specs_table.find_all('tr', class_='specifications-table_row')
            logger.info(f"Found {len(rows)} rows in {tab_id} ({tab_label})")
            for row in rows:
                cols = row.find_all('td', class_='specifications-table_col')
                logger.info(f"Processing row with {len(cols)} columns in {tab_id}")
                if len(cols) == 4:
                    key1 = cols[0].get_text(strip=True).rstrip(":")
                    value1 = cols[1].get_text(strip=True)
                    key2 = cols[2].get_text(strip=True).rstrip(":")
                    value2 = cols[3].get_text(strip=True)
                    specs_dict[key1] = value1
                    specs_dict[key2] = value2
                elif len(cols) == 2:
                    key = cols[0].get_text(strip=True).rstrip(":")
                    value = cols[1].get_text(strip=True)
                    specs_dict[key] = value
                
            logger.info(f"Extracted {len(specs_dict)} items from {tab_id} ({tab_label})")
            
            # Assign to the correct category based on label
            if "technical specifications" in tab_label:
                technical_specs.update(specs_dict)
                logger.info(f"Assigned {tab_id} as Technical Specifications with {len(specs_dict)} items")
            elif "general specifications" in tab_label:
                general_specs.update(specs_dict)
                logger.info(f"Assigned {tab_id} as General Specifications with {len(specs_dict)} items")
            else:
                logger.info(f"Skipping {tab_id} ({tab_label}) as it's not Technical or General Specifications")
        
        # Log final results
        logger.info(f"Scraped {len(technical_specs)} technical specifications")
        logger.info(f"Scraped {len(general_specs)} general specifications")
        
        # Check for overlap
        tech_keys = set(technical_specs.keys())
        gen_keys = set(general_specs.keys())
        overlap = tech_keys.intersection(gen_keys)
        if overlap:
            logger.warning(f"Overlap detected between technical and general specifications: {overlap}")
        
        return {
            "product_name": product_name,
            "key_features": key_features,
            "technical_specifications": technical_specs,
            "general_specifications": general_specs
        }
    except Exception as e:
        logger.error(f"Error in async scraping: {str(e)}")
        return {
            "product_name": "Unknown Product",
            "key_features": [],
            "technical_specifications": {},
            "general_specifications": {}
        }

async def async_search_confluence(query: str, session: aiohttp.ClientSession) -> str:
    """
    Asynchronously search Confluence for content matching the query.
    
    Args:
        query (str): Search query string
        session (aiohttp.ClientSession): Active aiohttp session
        
    Returns:
        str: Combined content from matching Confluence pages
    """
    if not session or session.closed:
        logger.error("Session is invalid or closed")
        return ""
        
    try:
        logger.info(f"Starting async Confluence search for query: '{query}'")
        
        # Build request parameters
        url = f"{CONFLUENCE_BASE_URL}/rest/api/content/search"
        normalized_query = normalize_text(query)
        cql_query = f'(text ~ "{normalized_query}") AND type = page'
        params = {
            "cql": cql_query,
            "expand": "body.storage,space,version",
            "limit": 10
        }
        auth = aiohttp.BasicAuth(login=CONFLUENCE_USERNAME, password=CONFLUENCE_API_TOKEN)
        
        # Add timeout and retry logic
        timeout = aiohttp.ClientTimeout(total=30)
        max_retries = 3
        retry_delay = 1
        
        for attempt in range(max_retries):
            try:
                async with session.get(
                    url, 
                    params=params, 
                    auth=auth, 
                    verify_ssl=False,
                    timeout=timeout
                ) as response:
                    if response.status == 429:  # Rate limit
                        retry_after = int(response.headers.get('Retry-After', retry_delay))
                        logger.warning(f"Rate limited, waiting {retry_after} seconds")
                        await asyncio.sleep(retry_after)
                        continue
                    response.raise_for_status()
                    results = await response.json()
                    break  # Success, exit retry loop
                    
            except asyncio.TimeoutError:
                if attempt == max_retries - 1:
                    logger.warning(f"Confluence API request timed out after {max_retries} attempts")
                    return ""
                logger.warning(f"Request timeout, attempt {attempt + 1}/{max_retries}")
                await asyncio.sleep(retry_delay)
                continue
                
            except aiohttp.ClientError as e:
                if attempt == max_retries - 1:
                    logger.warning(f"Confluence API request failed after {max_retries} attempts: {str(e)}")
                    return ""
                logger.warning(f"Request failed, attempt {attempt + 1}/{max_retries}: {str(e)}")
                await asyncio.sleep(retry_delay)
                continue
        
        # Process results
        pages = results.get("results", [])
        if not pages:
            logger.info(f"No Confluence pages found for query: {query}")
            return ""
            
        logger.info(f"Retrieved {len(pages)} pages from Confluence search")
        
        # Process pages in parallel
        tasks = []
        for page in pages:
            tasks.append(process_confluence_page(page))
        
        processed_contents = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out errors and combine content
        valid_contents = []
        for content in processed_contents:
            if isinstance(content, Exception):
                logger.error(f"Error processing page: {str(content)}")
                continue
            if content:
                valid_contents.append(content)
                
        combined_content = "\n\n".join(valid_contents)
        logger.info(f"Processed {len(valid_contents)} pages successfully")
        
        return combined_content
        
    except Exception as e:
        logger.error(f"Error in Confluence search: {str(e)}")
        return ""

async def process_confluence_page(page: dict) -> str:
    """Process a single Confluence page and extract relevant content."""
    try:
        page_title = page.get("title", "")
        page_space = page.get("space", {}).get("name", "")
        body = page.get("body", {}).get("storage", {}).get("value", "")
        
        if not body:
            logger.warning(f"No content found in page: {page_title}")
            return ""
            
        # Parse HTML content
        soup = BeautifulSoup(body, 'html.parser')
        
        # Remove unwanted elements
        for element in soup.find_all(['script', 'style', 'head']):
            element.decompose()
            
        # Extract text content
        text_content = soup.get_text(separator='\n', strip=True)
        
        # Format the content
        formatted_content = f"""
        Page: {page_title}
        Space: {page_space}
        Content:
        {text_content}
        """
        
        return formatted_content.strip()
        
    except Exception as e:
        logger.error(f"Error processing page {page.get('title', 'Unknown')}: {str(e)}")
        return ""

@app.get("/api/motor/sseusecase2/progress/{client_id}")
async def sse_progress(client_id: str):
    logger.info(f"Starting SSE for client_id: {client_id}")
    async def event_generator():
        while True:
            if client_id not in active_tasks:
                logger.info(f"Client {client_id} not in active_tasks, sending complete")
                yield {"event": "complete", "data": json.dumps({"message": "Task completed or disconnected", "percentage": 100})}
                break
            progress = active_tasks.get(client_id, {"message": "Waiting...", "percentage": 0})
            logger.info(f"Sending progress for {client_id}: {progress}")
            yield {"event": "progress", "data": json.dumps(progress)}
            await asyncio.sleep(1)
    return EventSourceResponse(event_generator())

# UseCase 3

# Global stores for progress updates and generated PDFs
progress_store: Dict[str, List[str]] = {}
pdf_store: Dict[str, BytesIO] = {}

###########################################
# UseCase 3: ReportLab Styles and DSPy Signatures
###########################################
STYLES = getSampleStyleSheet()
if 'Bullet' in STYLES.byName:
    bullet_style = STYLES.byName['Bullet']
    bullet_style.fontName = 'Helvetica'
    bullet_style.fontSize = 10
    bullet_style.leftIndent = 20
    bullet_style.bulletIndent = 10
    bullet_style.bulletFontName = 'Helvetica'
    bullet_style.bulletFontSize = 10
else:
    bullet_style = ParagraphStyle(
        name='Bullet',
        parent=STYLES['Normal'],
        leftIndent=20,
        bulletIndent=10,
        bulletFontName='Helvetica',
        bulletFontSize=10,
    )
    STYLES.add(bullet_style)

###########################################
# DSPy Signatures for Product Specification Content
###########################################
class GenerateProductSpecContent(Signature):
    """
    Default DSPy signature for generating content for a product specification section.
    """
    section_title: str = InputField(desc="Section title")
    product_category: str = InputField(desc="Product category (e.g., Motors, Bearings)")
    product_details: str = InputField(desc="Detailed product specifications")
    prompt: str = InputField(desc="Instructional prompt for content generation")
    output: str = OutputField(desc="Generated content for the section")

###########################################
# DSPy Signature for Product Manager Persona
###########################################
class GenerateProductSpecContentManager(Signature):
    """
    DSPy signature for generating content for a product specification section tailored for Product Managers.
    """
    section_title: str = InputField(desc="Section title")
    product_category: str = InputField(desc="Product category (e.g., Motors, Bearings)")
    product_details: str = InputField(desc="Detailed product specifications")
    prompt: str = InputField(desc="Instructional prompt for content generation for managers")
    output: str = OutputField(desc="Generated content for the section (Manager view)")

###########################################
# DSPy Signature for Product Engineer Persona
###########################################
class GenerateProductSpecContentEngineer(Signature):
    """
    DSPy signature for generating content for a product specification section tailored for Product Engineers.
    """
    section_title: str = InputField(desc="Section title")
    product_category: str = InputField(desc="Product category (e.g., Motors, Bearings)")
    product_details: str = InputField(desc="Detailed product specifications")
    prompt: str = InputField(desc="Instructional prompt for content generation for engineers")
    output: str = OutputField(desc="Generated content for the section (Engineer view)")

###########################################
# Helper: Add Decorations (Logo, Border, Footer)
###########################################
def add_decorations(canvas: reportlab_canvas.Canvas, doc: SimpleDocTemplate) -> None:  # type: ignore
    try:
        border_margin = 40.0
        extra_gap = 7.0
        page_width, page_height = doc.pagesize

        canvas.setLineWidth(1)
        canvas.setStrokeColor(colors.black)
        canvas.rect(border_margin, border_margin, page_width - 2 * border_margin, page_height - 2 * border_margin)

        logo_path = r".\ust-logo.png"
        logo = ImageReader(logo_path)
        orig_width, orig_height = logo.getSize()
        logo_width = (1.0 / 3.0) * inch
        aspect = orig_height / orig_width
        logo_height = logo_width * aspect
        logo_x = border_margin + extra_gap
        logo_y = page_height - border_margin - logo_height - extra_gap
        canvas.drawImage(logo_path, logo_x, logo_y, width=logo_width, height=logo_height, mask='auto', preserveAspectRatio=True)

        footer_text = "Fully AI generated and formatted"
        canvas.setFont("Helvetica", 8)
        text_width = canvas.stringWidth(footer_text, "Helvetica", 8)
        symbol_diameter = 15
        gap_between = 5
        total_width = symbol_diameter + gap_between + text_width
        group_right_x = page_width - border_margin
        group_start_x = group_right_x - total_width
        y_out = border_margin / 2

        symbol_radius = symbol_diameter / 2
        symbol_center_x = group_start_x + symbol_radius
        symbol_center_y = y_out + symbol_radius
        canvas.setFillColor(colors.red)
        canvas.circle(symbol_center_x, symbol_center_y, symbol_radius, stroke=0, fill=1)
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 12)
        canvas.drawCentredString(symbol_center_x, symbol_center_y - 4, "!")
        
        canvas.setFillColor(colors.black)
        canvas.setFont("Helvetica", 8)
        text_x = group_start_x + symbol_diameter + gap_between
        canvas.drawString(text_x, y_out + symbol_radius - 4, footer_text)
    except Exception as e:
        logging.error(f"Error in add_decorations: {str(e)}")

###########################################
# Helper: Convert Markdown Bold to HTML
###########################################
def convert_markdown_to_html(text: str) -> str:
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = text.replace('**', '')
    return text

###########################################
# Helper: Parse and Clean Section Content
###########################################
def parse_section_content(section_content: str) -> Tuple[str, List[str]]:
    section_content = re.sub(r'\[\[\s*#*\s*Key Points\s*#*\s*\]\]', 'Key Points:', section_content, flags=re.IGNORECASE)
    pattern = re.compile(r'Key Points\s*:', re.IGNORECASE)
    parts = pattern.split(section_content, maxsplit=1)
    if len(parts) > 1:
        detailed_text = parts[0].strip()
        bullet_block = parts[1].strip()
        bullets = []
        for line in bullet_block.splitlines():
            line = line.strip()
            if line:
                clean_line = re.sub(r'^\s*\d+\.\s*', '', line)
                clean_line = convert_markdown_to_html(clean_line)
                clean_line = re.sub(r'^[-•\s]+', '', clean_line)
                clean_line = clean_line.strip()
                if clean_line:
                    bullets.append(clean_line)
        return detailed_text, bullets
    return section_content, []

###########################################
# Helper: Format Detailed Text for PDF
###########################################
def format_detailed_text(detailed_text: str) -> List:
    flowables = []
    for line in detailed_text.split("\n"):
        line = line.strip()
        if not line:
            flowables.append(Spacer(1, 0.1 * inch))
            continue
        line = convert_markdown_to_html(line)
        if line.startswith("###"):
            text = line.lstrip("#").strip()
            flowables.append(Paragraph(text, STYLES['Heading3']))
        elif re.match(r'^\d+\.', line):
            numbered_style = ParagraphStyle('Numbered', parent=STYLES['Normal'], leftIndent=20)
            flowables.append(Paragraph(line, numbered_style))
        elif line.startswith("-"):
            text = line.lstrip("-").strip()
            flowables.append(Paragraph("• " + text, STYLES['Bullet']))
        else:
            flowables.append(Paragraph(line, STYLES['Normal']))
    return flowables

###########################################
# Helper: Build Prompts for Each Section
###########################################
def get_product_spec_prompts(product_category: str, product_details: str, custom_template: str = None) -> dict:
    prompts = {}
    if custom_template and custom_template.strip():
        sections = custom_template.splitlines()
        for i, line in enumerate(sections, start=1):
            line = line.strip()
            if not line:
                continue
            if ':' in line:
                heading, prompt_text = map(str.strip, line.split(':', 1))
            else:
                heading = line
                prompt_text = f"Provide detailed information on {heading}."
            heading = re.sub(r'^\d+\.\s*', '', heading)
            numbered_heading = f"{i}. {heading}"
            context = f"Product Category: {product_category}\nProduct Details: {product_details}\n\n"
            prompts[numbered_heading] = f"{context}Task: {prompt_text} Include a bullet summary under 'Key Points:' at the end."
    return prompts

###########################################
# Helper: Generate PDF Using ReportLab
###########################################
def generate_pdf_usecase3(product_category: str, content: dict, progress_callback=None) -> BytesIO:
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        elements = []

        # Title Page
        title_text = f"Product Specification for {product_category}"
        elements.append(Paragraph(title_text, STYLES['Title']))
        elements.append(Spacer(1, 0.5 * inch))

        # Table of Contents with proper wrapping
        from reportlab.lib.styles import ParagraphStyle
        toc_heading_style = ParagraphStyle(name="TOCHeading", fontSize=10, leading=12)
        toc_data = [
            [Paragraph("Section", toc_heading_style), Paragraph("Page", toc_heading_style)]
        ]
        page_num = 2  # Title/TOC occupies page 1
        for section in content.keys():
            toc_data.append([
                Paragraph(section.strip(), toc_heading_style),
                Paragraph(str(page_num), toc_heading_style)
            ])
            page_num += 1

        toc_table = Table(toc_data, colWidths=[380, 88])
        toc_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(toc_table)
        elements.append(PageBreak())

        # Sections
        for section, section_content in content.items():
            elements.append(Paragraph(section.strip(), STYLES['Heading2']))
            detailed_text, bullet_list = parse_section_content(section_content)
            flowables = format_detailed_text(detailed_text)
            elements.extend(flowables)

            if progress_callback:
                progress_callback(f"Section '{section}' formatted and added to PDF.")

            if bullet_list:
                elements.append(Paragraph("Key Points", STYLES['Heading3']))
                for bullet in bullet_list:
                    if not bullet.startswith("•"):
                        bullet = "• " + bullet
                    elements.append(Paragraph(bullet, STYLES['Bullet']))
                    elements.append(Spacer(1, 0.1 * inch))
            elements.append(PageBreak())

        doc.build(elements, onFirstPage=add_decorations, onLaterPages=add_decorations)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logging.error(f"Error generating PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")

###########################################
# Async Helper: Generate Section Content with Persona Handling
###########################################
async def generate_section_content(section: str, product_category: str, product_details: str, prompt: str, persona: str) -> tuple:
    persona_lower = persona.lower()
    if persona_lower in ["manager", "product_manager"]:
        dsp_signature = GenerateProductSpecContentManager
    elif persona_lower in ["engineer", "product_engineer"]:
        dsp_signature = GenerateProductSpecContentEngineer
    else:
        raise HTTPException(status_code=400, detail=f"Unknown persona '{persona}'. Choose either 'product_manager' or 'product_engineer'.")
    
    generate_spec = Predict(dsp_signature)
    result = await asyncio.to_thread(
        generate_spec,
        section_title=section,
        product_category=product_category,
        product_details=product_details,
        prompt=prompt,
        temperature=0.7,
        max_tokens=1000
    )
    content = result.output.strip() if result.output.strip() else "No content available."
    return section, content

###########################################
# Background Task: Generate PDF & Update Progress
###########################################
async def generate_pdf_task(job_id: str, product_category: str, product_details: str, custom_template: str, persona: str) -> None:
    def update_progress(msg: str) -> None:
        progress_store[job_id].append(msg)
    
    progress_store[job_id] = []
    update_progress("Starting PDF generation process...")
    
    update_progress("Building prompts from product information.")
    spec_prompts = get_product_spec_prompts(product_category, product_details, custom_template)
    update_progress("Prompts built successfully.")
    
    update_progress("Generating content for each section concurrently...")
    tasks = [
        generate_section_content(section, product_category, product_details, prompt, persona)
        for section, prompt in spec_prompts.items()
    ]
    results = await asyncio.gather(*tasks)
    generated_content = {section: content for section, content in results}
    for section, _ in results:
        update_progress(f"Content generated for section: {section}")
    
    update_progress("All sections generated. Now formatting PDF.")
    pdf_buffer = await asyncio.to_thread(generate_pdf_usecase3, product_category, generated_content, update_progress)
    update_progress("PDF generated successfully.")
    
    pdf_store[job_id] = pdf_buffer
    update_progress("PDF is ready for download. Use the download endpoint /api/motor/download/{job_id}.")

###########################################
# Async Helper: Extract Text from PDF
###########################################
async def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text from PDF: {str(e)}")

###########################################
# Endpoint: Initiate PDF Generation
###########################################
@app.post("/api/motor/generate-product-designer-pdf")
async def generate_product_designer_pdf(
    background_tasks: BackgroundTasks,
    product_category: str = Form(..., description="Product category (e.g., 'Motors', 'Bearings')"),
    product_details: str = Form(..., description="Product details (e.g., 'General Purpose Motor, 230V, 1140 RPM')"),
    custom_template: str = Form("", description="Optional custom template with separate headings (e.g., 'Heading: Custom text')"),
    template_file: UploadFile = File(None, description="Optional template file (.txt, .docx, or .pdf). If provided, overrides custom_template."),
    persona: str = Form(..., description="Persona type: 'product_manager' or 'product_engineer'")
):
    available_products = [product["product_name"] for product in products_data.get("products", [])]
    if product_category not in available_products:
        raise HTTPException(status_code=404, detail=f"Product category '{product_category}' not found in available products")

    extracted_text = custom_template
    if template_file:
        file_content_type = template_file.content_type
        file_bytes = await template_file.read()
        if file_content_type == "text/plain":
            extracted_text = file_bytes.decode("utf-8")
        elif file_content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(BytesIO(file_bytes))
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
        elif file_content_type == "application/pdf":
            try:
                pdf_reader = PdfReader(BytesIO(file_bytes))
                extracted_text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text
                if not extracted_text.strip():
                    raise HTTPException(status_code=400, detail="Could not extract text from the uploaded PDF.")
            except Exception as e:
                logging.error(f"Error extracting text from PDF: {str(e)}")
                raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a .txt, .docx, or .pdf file.")

    job_id = str(uuid.uuid4())
    background_tasks.add_task(generate_pdf_task, job_id, product_category, product_details, extracted_text, persona)
    return {
        "job_id": job_id,
        "message": "PDF generation started. Use the job_id to track progress via SSE at /api/motor/sse/progress/{job_id}."
    }

###########################################
# Endpoint: Get Available Products
###########################################
@app.get("/api/motor/product")
async def get_products():
    product_names = [product["product_name"] for product in products_data.get("products", [])]
    return JSONResponse(content={"products": product_names})

###########################################
# Endpoint: SSE for Progress Updates
###########################################
@app.get("/api/motor/sse/progress/{job_id}")
async def sse_progress(job_id: str):
    async def event_generator():
        if job_id not in progress_store:
            yield {"event": "error", "data": "Job not found"}
            return
        last_index = 0
        while True:
            messages = progress_store.get(job_id, [])
            if last_index < len(messages):
                for msg in messages[last_index:]:
                    yield {"event": "progress", "data": msg}
                last_index = len(messages)
            if messages and messages[-1].startswith("PDF is ready"):
                yield {"event": "complete", "data": messages[-1]}
                break
            await asyncio.sleep(1)
    return EventSourceResponse(event_generator())

###########################################
# Endpoint: Download Generated PDF
###########################################
@app.get("/api/motor/download/{job_id}")
async def download_pdf(job_id: str):
    if job_id in pdf_store:
        pdf_buffer = pdf_store[job_id]
        filename = f"product_spec_{job_id}.pdf"
        return StreamingResponse(
            pdf_buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    else:
        raise HTTPException(status_code=404, detail="PDF not found or not generated yet.")

###########################################
# Endpoint: Extract Text from Document Files
@app.post("/api/motor/extract-pdf-text")
async def extract_text(template_file: UploadFile = File(...)) -> Dict[str, Any]:
    """
    Accepts a file (PDF, DOCX/DOC, or TXT) via multipart/form-data and extracts text.
    For DOCX/DOC files, converts to PDF first, then extracts text.
    For PDFs, processes newlines and punctuation to form paragraphs.
    Each non-empty line (or paragraph) is returned as an entity.
    """
    content_type = template_file.content_type
    try:
        if content_type == "application/pdf":
            file_content = await template_file.read()
            extracted_text = await extract_text_from_pdf(file_content)
            # Replace multiple newlines with a space
            extracted_text = re.sub(r'\n+', ' ', extracted_text)
            # Insert newline after period (assuming end of sentence)
            extracted_text = re.sub(r'\.\s+', '.\n', extracted_text)
            entities = [line.strip() for line in extracted_text.splitlines() if line.strip()]
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            pdf_content = await convert_to_pdf(template_file)
            extracted_text = await extract_text_from_pdf(pdf_content)
            # Apply same processing as PDF for consistency
            extracted_text = re.sub(r'\n+', ' ', extracted_text)
            extracted_text = re.sub(r'\.\s+', '.\n', extracted_text)
            entities = [line.strip() for line in extracted_text.splitlines() if line.strip()]
        elif content_type == "text/plain":
            file_content = await template_file.read()
            extracted_text = file_content.decode("utf-8")
            entities = [line.strip() for line in extracted_text.splitlines() if line.strip()]
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Only PDF, DOCX, DOC, and TXT files are supported."
            )
        
        return {"extractedText": extracted_text, "entities": entities}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ---------------------------
# Replace Playwright imports with Selenium
# ---------------------------
# from playwright.sync_api import sync_playwright  # Remove this
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
import asyncio
import time

# Replace the Playwright scraping function with Selenium implementation
async def scrape_with_selenium(url, wait_time=5):
    """
    Scrape the given URL using Selenium instead of Playwright
    while maintaining the same data structure for PDF generation.
    """
    # Set up Chrome options
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    
    try:
        # Initialize the Chrome driver
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
        
        # Navigate to the page
        driver.get(url)
        
        # Wait for content to load
        time.sleep(wait_time)
        
        # Extract the data you need
        # Replace these selectors with the ones you actually need
        # Make sure the data structure matches what your PDF generator expects
        data = {}
        
        # Example: Extract title
        try:
            h1_elements = driver.find_elements(By.TAG_NAME, 'h1')
            if h1_elements:
                data['title'] = h1_elements[0].text
        except Exception as e:
            print(f"Error extracting title: {e}")
        
        # Example: Extract paragraphs
        try:
            paragraphs = driver.find_elements(By.TAG_NAME, 'p')
            data['content'] = [p.text for p in paragraphs]
        except Exception as e:
            print(f"Error extracting paragraphs: {e}")
            data['content'] = []
        
        # Example: Extract images
        try:
            images = driver.find_elements(By.TAG_NAME, 'img')
            data['images'] = [img.get_attribute('src') for img in images]
        except Exception as e:
            print(f"Error extracting images: {e}")
            data['images'] = []
        
        # Add any other data extraction you need
        # ...
        
        # Take a screenshot if needed
        # driver.save_screenshot("screenshot.png")
        
        return data
    except Exception as e:
        print(f"Error scraping with Selenium: {e}")
        return None
    finally:
        if 'driver' in locals():
            driver.quit()

# Update the endpoint to use the new Selenium scraping function
@app.route('/scrape', methods=['POST'])
def scrape():
    data = request.json
    url = data.get('url')
    
    if not url:
        return jsonify({'error': 'URL is required'}), 400
    
    try:
        # Use the new Selenium scraping function instead of Playwright
        scraped_data = asyncio.run(scrape_with_selenium(url))
        
        if not scraped_data:
            return jsonify({'error': 'Failed to scrape the URL'}), 500
        
        # The rest of your code for PDF generation should remain the same
        # as the data structure from the scraping function matches the expected format
        
        # Generate PDF with the scraped data
        pdf_path = generate_pdf(scraped_data)  # Assuming this function exists
        
        # Return the PDF or a download link
        return jsonify({'success': True, 'pdf_path': pdf_path})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app)